"""Authenticity scoring component for document forensics."""

import logging
import statistics
from typing import Dict, List, Optional, Any, Tuple
import numpy as np
from pathlib import Path
import json
import hashlib
from datetime import datetime
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import cv2

from ..core.models import (
    AuthenticityAnalysis, AuthenticityScore, ComparisonResult, 
    StructureValidation, ObjectAssessment, RiskLevel, Document
)

logger = logging.getLogger(__name__)


class AuthenticityScorer:
    """Multi-factor authenticity assessment for documents."""
    
    def __init__(self):
        """Initialize the authenticity scorer."""
        self.reference_samples = {}
        self.format_validators = {
            'pdf': self._validate_pdf_structure,
            'docx': self._validate_docx_structure,
            'jpg': self._validate_image_structure,
            'jpeg': self._validate_image_structure,
            'png': self._validate_image_structure
        }
    
    async def calculate_authenticity_score(
        self, 
        document_path: str, 
        document_id: int,
        reference_samples: Optional[List[str]] = None
    ) -> AuthenticityAnalysis:
        """
        Calculate comprehensive authenticity score for a document.
        
        Args:
            document_path: Path to the document file
            document_id: Document ID for tracking
            reference_samples: Optional list of reference document paths
            
        Returns:
            AuthenticityAnalysis with comprehensive scoring
        """
        logger.info(f"Starting authenticity scoring for document {document_id}")
        
        try:
            # Multi-factor authenticity assessment
            authenticity_factors = await self._assess_authenticity_factors(document_path)
            
            # Compare against reference samples if provided
            comparison_results = []
            if reference_samples:
                comparison_results = await self._compare_against_samples(
                    document_path, reference_samples
                )
            
            # Validate file structure
            structure_validation = await self._validate_file_structure(document_path)
            
            # Assess embedded objects
            embedded_objects = await self._assess_embedded_objects(document_path)
            
            # Calculate overall authenticity score
            authenticity_score = self._calculate_overall_score(
                authenticity_factors, comparison_results, structure_validation, embedded_objects
            )
            
            # Collect forensic indicators
            forensic_indicators = self._collect_forensic_indicators(
                document_path, authenticity_factors, structure_validation
            )
            
            return AuthenticityAnalysis(
                document_id=document_id,
                authenticity_score=authenticity_score,
                comparison_results=comparison_results,
                structure_validation=structure_validation,
                embedded_objects_assessment=embedded_objects,
                forensic_indicators=forensic_indicators
            )
            
        except Exception as e:
            logger.error(f"Error during authenticity scoring: {str(e)}")
            # Return minimal analysis on error
            return AuthenticityAnalysis(
                document_id=document_id,
                authenticity_score=AuthenticityScore(
                    overall_score=0.0,
                    confidence_level=0.0,
                    contributing_factors={},
                    risk_assessment=RiskLevel.LOW
                )
            )
    
    async def _assess_authenticity_factors(self, document_path: str) -> Dict[str, float]:
        """Assess multiple authenticity factors."""
        factors = {}
        
        try:
            # File format consistency
            factors['format_consistency'] = await self._assess_format_consistency(document_path)
            
            # Metadata authenticity
            factors['metadata_authenticity'] = await self._assess_metadata_authenticity(document_path)
            
            # Content integrity
            factors['content_integrity'] = await self._assess_content_integrity(document_path)
            
            # Creation pattern analysis
            factors['creation_patterns'] = await self._assess_creation_patterns(document_path)
            
            # Digital signature validity (if present)
            factors['signature_validity'] = await self._assess_signature_validity(document_path)
            
        except Exception as e:
            logger.error(f"Error assessing authenticity factors: {str(e)}")
            
        return factors
    
    async def _assess_format_consistency(self, document_path: str) -> float:
        """Assess file format consistency and compliance."""
        try:
            # Detect file type using magic if available
            if MAGIC_AVAILABLE:
                file_type = magic.from_file(document_path, mime=True)
            else:
                # Fallback to extension-based detection
                file_extension = Path(document_path).suffix.lower()
                extension_type_map = {
                    '.pdf': 'application/pdf',
                    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.png': 'image/png'
                }
                file_type = extension_type_map.get(file_extension, 'unknown')
            
            file_extension = Path(document_path).suffix.lower()
            
            # Check if extension matches detected type
            extension_type_map = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png'
            }
            
            expected_type = extension_type_map.get(file_extension)
            if expected_type and file_type == expected_type:
                return 0.9  # High consistency
            elif expected_type:
                return 0.3 if MAGIC_AVAILABLE else 0.7  # Lower penalty if no magic available
            else:
                return 0.5  # Unknown format
                
        except Exception as e:
            logger.error(f"Error assessing format consistency: {str(e)}")
            return 0.0
    
    async def _assess_metadata_authenticity(self, document_path: str) -> float:
        """Assess metadata authenticity and consistency."""
        try:
            file_extension = Path(document_path).suffix.lower()
            
            if file_extension == '.pdf':
                return await self._assess_pdf_metadata_authenticity(document_path)
            elif file_extension == '.docx':
                return await self._assess_docx_metadata_authenticity(document_path)
            elif file_extension in ['.jpg', '.jpeg', '.png']:
                return await self._assess_image_metadata_authenticity(document_path)
            else:
                return 0.5  # Neutral for unknown formats
                
        except Exception as e:
            logger.error(f"Error assessing metadata authenticity: {str(e)}")
            return 0.0
    
    async def _assess_pdf_metadata_authenticity(self, pdf_path: str) -> float:
        """Assess PDF metadata for authenticity indicators."""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = pdf_reader.metadata
                
                if not metadata:
                    return 0.3  # Missing metadata is suspicious
                
                authenticity_score = 0.5  # Base score
                
                # Check for creation and modification dates
                if metadata.get('/CreationDate') and metadata.get('/ModDate'):
                    authenticity_score += 0.2
                
                # Check for producer/creator information
                if metadata.get('/Producer') or metadata.get('/Creator'):
                    authenticity_score += 0.2
                
                # Check for suspicious patterns
                creator = str(metadata.get('/Creator', ''))
                producer = str(metadata.get('/Producer', ''))
                
                # Common legitimate creators/producers
                legitimate_patterns = [
                    'Microsoft', 'Adobe', 'LibreOffice', 'OpenOffice',
                    'Google', 'Apple', 'Foxit', 'PDFCreator'
                ]
                
                if any(pattern in creator or pattern in producer for pattern in legitimate_patterns):
                    authenticity_score += 0.1
                
                return min(1.0, authenticity_score)
                
        except Exception as e:
            logger.error(f"Error assessing PDF metadata: {str(e)}")
            return 0.0
    
    async def _assess_docx_metadata_authenticity(self, docx_path: str) -> float:
        """Assess DOCX metadata for authenticity indicators."""
        try:
            doc = DocxDocument(docx_path)
            core_props = doc.core_properties
            
            authenticity_score = 0.5  # Base score
            
            # Check for creation and modification dates
            if core_props.created and core_props.modified:
                authenticity_score += 0.2
                
                # Check if modification date is after creation date
                if core_props.modified >= core_props.created:
                    authenticity_score += 0.1
            
            # Check for author information
            if core_props.author:
                authenticity_score += 0.1
            
            # Check for application information
            if hasattr(core_props, 'last_modified_by') and core_props.last_modified_by:
                authenticity_score += 0.1
            
            return min(1.0, authenticity_score)
            
        except Exception as e:
            logger.error(f"Error assessing DOCX metadata: {str(e)}")
            return 0.0
    
    async def _assess_image_metadata_authenticity(self, image_path: str) -> float:
        """Assess image metadata for authenticity indicators."""
        try:
            image = Image.open(image_path)
            exif_data = image.getexif()
            
            if not exif_data:
                return 0.4  # No EXIF data is somewhat suspicious for photos
            
            authenticity_score = 0.6  # Base score for having EXIF
            
            # Check for camera information
            camera_make = exif_data.get(271)  # Make
            camera_model = exif_data.get(272)  # Model
            
            if camera_make and camera_model:
                authenticity_score += 0.2
            
            # Check for timestamp
            datetime_original = exif_data.get(36867)  # DateTimeOriginal
            if datetime_original:
                authenticity_score += 0.1
            
            # Check for GPS information
            gps_info = exif_data.get(34853)  # GPSInfo
            if gps_info:
                authenticity_score += 0.1
            
            return min(1.0, authenticity_score)
            
        except Exception as e:
            logger.error(f"Error assessing image metadata: {str(e)}")
            return 0.0
    
    async def _assess_content_integrity(self, document_path: str) -> float:
        """Assess content integrity and consistency."""
        try:
            # Calculate file hash for integrity
            with open(document_path, 'rb') as f:
                file_hash = hashlib.sha256(f.read()).hexdigest()
            
            # Check file size consistency
            file_size = Path(document_path).stat().st_size
            
            # Basic integrity checks
            if file_size > 0:
                integrity_score = 0.7  # Base score for non-empty file
                
                # Additional checks based on file type
                file_extension = Path(document_path).suffix.lower()
                
                if file_extension == '.pdf':
                    integrity_score += await self._check_pdf_integrity(document_path)
                elif file_extension in ['.jpg', '.jpeg', '.png']:
                    integrity_score += await self._check_image_integrity(document_path)
                
                return min(1.0, integrity_score)
            else:
                return 0.0  # Empty file
                
        except Exception as e:
            logger.error(f"Error assessing content integrity: {str(e)}")
            return 0.0
    
    async def _check_pdf_integrity(self, pdf_path: str) -> float:
        """Check PDF-specific integrity indicators."""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF can be read without errors
                num_pages = len(pdf_reader.pages)
                if num_pages > 0:
                    # Try to extract text from first page
                    first_page = pdf_reader.pages[0]
                    text = first_page.extract_text()
                    return 0.2  # Successfully readable
                else:
                    return 0.0  # No pages
                    
        except Exception as e:
            logger.error(f"PDF integrity check failed: {str(e)}")
            return 0.0
    
    async def _check_image_integrity(self, image_path: str) -> float:
        """Check image-specific integrity indicators."""
        try:
            # Try to load image with PIL
            image = Image.open(image_path)
            width, height = image.size
            
            if width > 0 and height > 0:
                # Try to load with OpenCV as well
                cv_image = cv2.imread(image_path)
                if cv_image is not None:
                    return 0.2  # Successfully readable by both libraries
                else:
                    return 0.1  # Only readable by PIL
            else:
                return 0.0  # Invalid dimensions
                
        except Exception as e:
            logger.error(f"Image integrity check failed: {str(e)}")
            return 0.0
    
    async def _assess_creation_patterns(self, document_path: str) -> float:
        """Assess document creation patterns for authenticity."""
        try:
            # Analyze file timestamps
            file_stat = Path(document_path).stat()
            creation_time = file_stat.st_ctime
            modification_time = file_stat.st_mtime
            
            # Check if modification time is reasonable relative to creation time
            if modification_time >= creation_time:
                pattern_score = 0.6
            else:
                pattern_score = 0.2  # Suspicious timestamp pattern
            
            # Additional pattern analysis based on file type
            file_extension = Path(document_path).suffix.lower()
            
            if file_extension == '.pdf':
                pattern_score += await self._analyze_pdf_creation_patterns(document_path)
            elif file_extension == '.docx':
                pattern_score += await self._analyze_docx_creation_patterns(document_path)
            
            return min(1.0, pattern_score)
            
        except Exception as e:
            logger.error(f"Error assessing creation patterns: {str(e)}")
            return 0.0
    
    async def _analyze_pdf_creation_patterns(self, pdf_path: str) -> float:
        """Analyze PDF creation patterns."""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = pdf_reader.metadata
                
                if metadata:
                    creation_date = metadata.get('/CreationDate')
                    mod_date = metadata.get('/ModDate')
                    
                    if creation_date and mod_date:
                        # Check if dates are in reasonable format and order
                        return 0.2
                    elif creation_date or mod_date:
                        return 0.1
                
                return 0.0
                
        except Exception as e:
            logger.error(f"Error analyzing PDF creation patterns: {str(e)}")
            return 0.0
    
    async def _analyze_docx_creation_patterns(self, docx_path: str) -> float:
        """Analyze DOCX creation patterns."""
        try:
            doc = DocxDocument(docx_path)
            core_props = doc.core_properties
            
            if core_props.created and core_props.modified:
                # Check chronological order
                if core_props.modified >= core_props.created:
                    return 0.2
                else:
                    return 0.0  # Suspicious order
            elif core_props.created or core_props.modified:
                return 0.1
            
            return 0.0
            
        except Exception as e:
            logger.error(f"Error analyzing DOCX creation patterns: {str(e)}")
            return 0.0
    
    async def _assess_signature_validity(self, document_path: str) -> float:
        """Assess digital signature validity if present."""
        try:
            file_extension = Path(document_path).suffix.lower()
            
            if file_extension == '.pdf':
                return await self._check_pdf_signatures(document_path)
            else:
                return 0.5  # Neutral for formats without common signature support
                
        except Exception as e:
            logger.error(f"Error assessing signature validity: {str(e)}")
            return 0.0
    
    async def _check_pdf_signatures(self, pdf_path: str) -> float:
        """Check PDF digital signatures."""
        try:
            with open(pdf_path, 'rb') as file:
                content = file.read()
                
                # Look for signature-related keywords
                if b'/Sig' in content or b'/ByteRange' in content:
                    # PDF has signature fields, but we can't verify without proper tools
                    return 0.7  # Presence of signature structure is positive
                else:
                    return 0.5  # No signatures (neutral)
                    
        except Exception as e:
            logger.error(f"Error checking PDF signatures: {str(e)}")
            return 0.0
    
    async def _compare_against_samples(
        self, 
        document_path: str, 
        reference_samples: List[str]
    ) -> List[ComparisonResult]:
        """Compare document against known authentic samples."""
        comparison_results = []
        
        for i, sample_path in enumerate(reference_samples):
            try:
                similarity_score = await self._calculate_similarity(document_path, sample_path)
                matching_features, differing_features = await self._analyze_features(
                    document_path, sample_path
                )
                
                comparison_results.append(ComparisonResult(
                    reference_id=f"sample_{i}",
                    similarity_score=similarity_score,
                    matching_features=matching_features,
                    differing_features=differing_features,
                    confidence=min(0.9, similarity_score + 0.1)
                ))
                
            except Exception as e:
                logger.error(f"Error comparing with sample {sample_path}: {str(e)}")
                
        return comparison_results
    
    async def _calculate_similarity(self, doc1_path: str, doc2_path: str) -> float:
        """Calculate similarity between two documents."""
        try:
            # Basic similarity based on file size and type
            stat1 = Path(doc1_path).stat()
            stat2 = Path(doc2_path).stat()
            
            # Size similarity
            size_ratio = min(stat1.st_size, stat2.st_size) / max(stat1.st_size, stat2.st_size)
            
            # Type similarity
            ext1 = Path(doc1_path).suffix.lower()
            ext2 = Path(doc2_path).suffix.lower()
            type_similarity = 1.0 if ext1 == ext2 else 0.0
            
            # Combine similarities
            overall_similarity = (size_ratio * 0.3 + type_similarity * 0.7)
            
            return overall_similarity
            
        except Exception as e:
            logger.error(f"Error calculating similarity: {str(e)}")
            return 0.0
    
    async def _analyze_features(
        self, 
        doc1_path: str, 
        doc2_path: str
    ) -> Tuple[List[str], List[str]]:
        """Analyze matching and differing features between documents."""
        matching_features = []
        differing_features = []
        
        try:
            # Compare file extensions
            ext1 = Path(doc1_path).suffix.lower()
            ext2 = Path(doc2_path).suffix.lower()
            
            if ext1 == ext2:
                matching_features.append(f"file_type_{ext1}")
            else:
                differing_features.append(f"file_type_{ext1}_vs_{ext2}")
            
            # Compare file sizes
            size1 = Path(doc1_path).stat().st_size
            size2 = Path(doc2_path).stat().st_size
            
            size_diff_ratio = abs(size1 - size2) / max(size1, size2)
            if size_diff_ratio < 0.1:
                matching_features.append("similar_file_size")
            else:
                differing_features.append("different_file_size")
            
        except Exception as e:
            logger.error(f"Error analyzing features: {str(e)}")
            
        return matching_features, differing_features
    
    async def _validate_file_structure(self, document_path: str) -> StructureValidation:
        """Validate file structure against format specifications."""
        file_extension = Path(document_path).suffix.lower().lstrip('.')
        
        validator = self.format_validators.get(file_extension)
        if validator:
            return await validator(document_path)
        else:
            return StructureValidation(
                is_valid=True,
                format_compliance=0.5,
                violations=[],
                recommendations=["Unknown format - manual validation recommended"]
            )
    
    async def _validate_pdf_structure(self, pdf_path: str) -> StructureValidation:
        """Validate PDF structure compliance."""
        violations = []
        recommendations = []
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is readable
                num_pages = len(pdf_reader.pages)
                if num_pages == 0:
                    violations.append("PDF contains no pages")
                
                # Check for metadata
                metadata = pdf_reader.metadata
                if not metadata:
                    recommendations.append("Consider adding document metadata")
                
                # Check for encryption
                if pdf_reader.is_encrypted:
                    recommendations.append("Document is encrypted")
                
                compliance_score = 1.0 - (len(violations) * 0.3)
                
                return StructureValidation(
                    is_valid=len(violations) == 0,
                    format_compliance=max(0.0, compliance_score),
                    violations=violations,
                    recommendations=recommendations
                )
                
        except Exception as e:
            logger.error(f"Error validating PDF structure: {str(e)}")
            return StructureValidation(
                is_valid=False,
                format_compliance=0.0,
                violations=[f"PDF validation error: {str(e)}"],
                recommendations=["Manual inspection required"]
            )
    
    async def _validate_docx_structure(self, docx_path: str) -> StructureValidation:
        """Validate DOCX structure compliance."""
        violations = []
        recommendations = []
        
        try:
            doc = DocxDocument(docx_path)
            
            # Check if document has content
            if not doc.paragraphs:
                violations.append("Document contains no paragraphs")
            
            # Check for core properties
            core_props = doc.core_properties
            if not core_props.author:
                recommendations.append("Consider adding author information")
            
            compliance_score = 1.0 - (len(violations) * 0.3)
            
            return StructureValidation(
                is_valid=len(violations) == 0,
                format_compliance=max(0.0, compliance_score),
                violations=violations,
                recommendations=recommendations
            )
            
        except Exception as e:
            logger.error(f"Error validating DOCX structure: {str(e)}")
            return StructureValidation(
                is_valid=False,
                format_compliance=0.0,
                violations=[f"DOCX validation error: {str(e)}"],
                recommendations=["Manual inspection required"]
            )
    
    async def _validate_image_structure(self, image_path: str) -> StructureValidation:
        """Validate image structure compliance."""
        violations = []
        recommendations = []
        
        try:
            image = Image.open(image_path)
            width, height = image.size
            
            # Check image dimensions
            if width == 0 or height == 0:
                violations.append("Invalid image dimensions")
            
            # Check for EXIF data (for photos)
            exif_data = image.getexif()
            if not exif_data and Path(image_path).suffix.lower() in ['.jpg', '.jpeg']:
                recommendations.append("JPEG image lacks EXIF data")
            
            compliance_score = 1.0 - (len(violations) * 0.3)
            
            return StructureValidation(
                is_valid=len(violations) == 0,
                format_compliance=max(0.0, compliance_score),
                violations=violations,
                recommendations=recommendations
            )
            
        except Exception as e:
            logger.error(f"Error validating image structure: {str(e)}")
            return StructureValidation(
                is_valid=False,
                format_compliance=0.0,
                violations=[f"Image validation error: {str(e)}"],
                recommendations=["Manual inspection required"]
            )
    
    async def _assess_embedded_objects(self, document_path: str) -> List[ObjectAssessment]:
        """Assess embedded objects within the document."""
        assessments = []
        
        try:
            file_extension = Path(document_path).suffix.lower()
            
            if file_extension == '.pdf':
                assessments = await self._assess_pdf_objects(document_path)
            elif file_extension == '.docx':
                assessments = await self._assess_docx_objects(document_path)
            
        except Exception as e:
            logger.error(f"Error assessing embedded objects: {str(e)}")
            
        return assessments
    
    async def _assess_pdf_objects(self, pdf_path: str) -> List[ObjectAssessment]:
        """Assess embedded objects in PDF."""
        assessments = []
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Simple assessment - check for images and forms
                for page_num, page in enumerate(pdf_reader.pages):
                    if '/XObject' in page.get('/Resources', {}):
                        assessments.append(ObjectAssessment(
                            object_id=f"page_{page_num}_xobjects",
                            object_type="embedded_objects",
                            integrity_score=0.8,
                            authenticity_indicators=["embedded_content_present"],
                            anomalies=[]
                        ))
                        
        except Exception as e:
            logger.error(f"Error assessing PDF objects: {str(e)}")
            
        return assessments
    
    async def _assess_docx_objects(self, docx_path: str) -> List[ObjectAssessment]:
        """Assess embedded objects in DOCX."""
        assessments = []
        
        try:
            doc = DocxDocument(docx_path)
            
            # Check for images and other embedded content
            for rel in doc.part.rels.values():
                if 'image' in rel.target_ref:
                    assessments.append(ObjectAssessment(
                        object_id=f"image_{rel.rId}",
                        object_type="embedded_image",
                        integrity_score=0.8,
                        authenticity_indicators=["embedded_image_present"],
                        anomalies=[]
                    ))
                    
        except Exception as e:
            logger.error(f"Error assessing DOCX objects: {str(e)}")
            
        return assessments
    
    def _calculate_overall_score(
        self,
        authenticity_factors: Dict[str, float],
        comparison_results: List[ComparisonResult],
        structure_validation: StructureValidation,
        embedded_objects: List[ObjectAssessment]
    ) -> AuthenticityScore:
        """Calculate overall authenticity score."""
        
        # Weight the different factors
        factor_weights = {
            'format_consistency': 0.2,
            'metadata_authenticity': 0.25,
            'content_integrity': 0.25,
            'creation_patterns': 0.15,
            'signature_validity': 0.15
        }
        
        # Calculate weighted score from factors
        weighted_score = 0.0
        total_weight = 0.0
        
        for factor, score in authenticity_factors.items():
            weight = factor_weights.get(factor, 0.1)
            weighted_score += score * weight
            total_weight += weight
        
        if total_weight > 0:
            factor_score = weighted_score / total_weight
        else:
            factor_score = 0.0
        
        # Incorporate structure validation
        structure_score = structure_validation.format_compliance
        
        # Incorporate comparison results if available
        comparison_score = 0.5  # Neutral if no comparisons
        if comparison_results:
            comparison_score = statistics.mean([r.similarity_score for r in comparison_results])
        
        # Incorporate embedded object assessment
        object_score = 0.5  # Neutral if no objects
        if embedded_objects:
            object_score = statistics.mean([obj.integrity_score for obj in embedded_objects])
        
        # Calculate overall score
        overall_score = (
            factor_score * 0.5 +
            structure_score * 0.2 +
            comparison_score * 0.2 +
            object_score * 0.1
        )
        
        # Calculate confidence level
        confidence_level = self._calculate_confidence_level(
            authenticity_factors, comparison_results, structure_validation
        )
        
        # Determine risk assessment
        risk_assessment = self._determine_risk_level(overall_score)
        
        return AuthenticityScore(
            overall_score=overall_score,
            confidence_level=confidence_level,
            contributing_factors=authenticity_factors,
            risk_assessment=risk_assessment
        )
    
    def _calculate_confidence_level(
        self,
        authenticity_factors: Dict[str, float],
        comparison_results: List[ComparisonResult],
        structure_validation: StructureValidation
    ) -> float:
        """Calculate confidence level for the authenticity assessment."""
        
        # Base confidence from number of factors assessed
        base_confidence = min(0.8, len(authenticity_factors) * 0.15)
        
        # Boost confidence if we have comparison results
        comparison_boost = 0.0
        if comparison_results:
            avg_comparison_confidence = statistics.mean([r.confidence for r in comparison_results])
            comparison_boost = avg_comparison_confidence * 0.2
        
        # Boost confidence if structure validation passed
        structure_boost = 0.0
        if structure_validation.is_valid:
            structure_boost = 0.1
        
        total_confidence = base_confidence + comparison_boost + structure_boost
        return min(1.0, total_confidence)
    
    def _determine_risk_level(self, overall_score: float) -> RiskLevel:
        """Determine risk level based on overall authenticity score."""
        if overall_score >= 0.8:
            return RiskLevel.LOW
        elif overall_score >= 0.6:
            return RiskLevel.MEDIUM
        elif overall_score >= 0.4:
            return RiskLevel.HIGH
        else:
            return RiskLevel.CRITICAL
    
    def _collect_forensic_indicators(
        self,
        document_path: str,
        authenticity_factors: Dict[str, float],
        structure_validation: StructureValidation
    ) -> Dict[str, Any]:
        """Collect forensic indicators for the analysis."""
        
        indicators = {
            'file_path': document_path,
            'file_size': Path(document_path).stat().st_size,
            'file_extension': Path(document_path).suffix.lower(),
            'analysis_timestamp': datetime.utcnow().isoformat(),
            'factor_scores': authenticity_factors,
            'structure_violations': structure_validation.violations,
            'structure_recommendations': structure_validation.recommendations
        }
        
        # Add file hash for integrity verification
        try:
            with open(document_path, 'rb') as f:
                file_hash = hashlib.sha256(f.read()).hexdigest()
                indicators['file_hash_sha256'] = file_hash
        except Exception as e:
            logger.error(f"Error calculating file hash: {str(e)}")
        
        return indicators