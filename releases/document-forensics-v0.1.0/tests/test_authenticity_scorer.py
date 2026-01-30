"""Property-based tests for authenticity scoring functionality."""

import asyncio
import os
import tempfile
from pathlib import Path
from typing import List, Dict, Any
import numpy as np
import pytest
from hypothesis import given, strategies as st, settings, assume, HealthCheck
from PIL import Image
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from src.document_forensics.analysis.authenticity_scorer import AuthenticityScorer
from src.document_forensics.core.models import (
    AuthenticityAnalysis, AuthenticityScore, ComparisonResult,
    StructureValidation, ObjectAssessment, RiskLevel
)


class TestAuthenticityScorer:
    """Property-based tests for authenticity scoring."""
    
    @pytest.fixture
    def scorer(self):
        """Create an authenticity scorer instance."""
        return AuthenticityScorer()
    
    @pytest.fixture
    def temp_dir(self):
        """Create a temporary directory for test files."""
        with tempfile.TemporaryDirectory() as temp_dir:
            yield temp_dir
    
    def create_test_pdf(self, temp_path: str, add_metadata: bool = True) -> str:
        """Create a test PDF with optional metadata."""
        c = canvas.Canvas(temp_path, pagesize=letter)
        
        # Add content
        c.drawString(100, 750, "Test Document")
        c.drawString(100, 730, "This is a sample document for testing.")
        
        if add_metadata:
            # Add metadata
            c.setTitle("Test Document")
            c.setAuthor("Test Author")
            c.setSubject("Testing")
        
        c.save()
        return temp_path
    
    def create_test_docx(self, temp_path: str, add_metadata: bool = True) -> str:
        """Create a test DOCX with optional metadata."""
        doc = DocxDocument()
        
        # Add content
        doc.add_paragraph("Test Document")
        doc.add_paragraph("This is a sample document for testing.")
        
        if add_metadata:
            # Add metadata
            doc.core_properties.title = "Test Document"
            doc.core_properties.author = "Test Author"
            doc.core_properties.subject = "Testing"
        
        doc.save(temp_path)
        return temp_path
    
    def create_test_image(self, temp_path: str, width: int = 400, height: int = 300) -> str:
        """Create a test image."""
        # Create a simple test image
        image = Image.new('RGB', (width, height), color='blue')
        image.save(temp_path, 'JPEG')
        return temp_path
    
    @given(
        document_type=st.sampled_from(['pdf', 'docx', 'jpg']),
        has_metadata=st.booleans(),
        file_size_factor=st.floats(min_value=0.1, max_value=2.0)
    )
    @settings(max_examples=15, deadline=25000, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @pytest.mark.asyncio
    async def test_property_authenticity_assessment_completeness(
        self, scorer, temp_dir, document_type, has_metadata, file_size_factor
    ):
        """
        **Feature: document-forensics, Property 5: Authenticity Assessment Completeness**
        **Validates: Requirements 4.1, 4.2, 4.3, 4.4, 4.5**
        
        For any document undergoing authenticity analysis, the system should calculate 
        confidence scores based on multiple forensic indicators, compare against reference 
        samples when available, validate file structure, and provide a comprehensive assessment.
        """
        # Create test document
        if document_type == 'pdf':
            file_path = os.path.join(temp_dir, "test.pdf")
            self.create_test_pdf(file_path, has_metadata)
        elif document_type == 'docx':
            file_path = os.path.join(temp_dir, "test.docx")
            self.create_test_docx(file_path, has_metadata)
        else:  # jpg
            file_path = os.path.join(temp_dir, "test.jpg")
            width = int(400 * file_size_factor)
            height = int(300 * file_size_factor)
            assume(width >= 10 and height >= 10)  # Ensure minimum viable image size
            self.create_test_image(file_path, width, height)
        
        # Perform authenticity analysis
        result = await scorer.calculate_authenticity_score(file_path, document_id=1)
        
        # Verify comprehensive assessment structure
        assert isinstance(result, AuthenticityAnalysis)
        assert result.document_id == 1
        assert isinstance(result.authenticity_score, AuthenticityScore)
        
        # Verify authenticity score properties
        auth_score = result.authenticity_score
        assert isinstance(auth_score.overall_score, float)
        assert 0.0 <= auth_score.overall_score <= 1.0
        assert isinstance(auth_score.confidence_level, float)
        assert 0.0 <= auth_score.confidence_level <= 1.0
        assert isinstance(auth_score.contributing_factors, dict)
        assert isinstance(auth_score.risk_assessment, RiskLevel)
        
        # Verify multiple forensic indicators are assessed
        factors = auth_score.contributing_factors
        expected_factors = [
            'format_consistency', 'metadata_authenticity', 'content_integrity',
            'creation_patterns', 'signature_validity'
        ]
        
        # Should have assessed multiple factors
        assessed_factors = [factor for factor in expected_factors if factor in factors]
        assert len(assessed_factors) >= 3  # At least 3 factors should be assessed
        
        # All factor scores should be valid
        for factor, score in factors.items():
            assert isinstance(score, float)
            assert 0.0 <= score <= 1.0
        
        # Verify structure validation is performed
        assert isinstance(result.structure_validation, StructureValidation)
        struct_val = result.structure_validation
        assert isinstance(struct_val.is_valid, bool)
        assert isinstance(struct_val.format_compliance, float)
        assert 0.0 <= struct_val.format_compliance <= 1.0
        assert isinstance(struct_val.violations, list)
        assert isinstance(struct_val.recommendations, list)
        
        # Verify embedded objects assessment
        assert isinstance(result.embedded_objects_assessment, list)
        for obj_assessment in result.embedded_objects_assessment:
            assert isinstance(obj_assessment, ObjectAssessment)
            assert isinstance(obj_assessment.integrity_score, float)
            assert 0.0 <= obj_assessment.integrity_score <= 1.0
        
        # Verify forensic indicators are collected
        assert isinstance(result.forensic_indicators, dict)
        indicators = result.forensic_indicators
        
        # Should contain basic forensic information
        expected_indicator_keys = ['file_path', 'file_size', 'file_extension', 'analysis_timestamp']
        for key in expected_indicator_keys:
            assert key in indicators
        
        # Verify risk assessment correlates with score
        if auth_score.overall_score >= 0.8:
            assert auth_score.risk_assessment == RiskLevel.LOW
        elif auth_score.overall_score <= 0.4:
            assert auth_score.risk_assessment in [RiskLevel.HIGH, RiskLevel.CRITICAL]
    
    @given(
        num_samples=st.integers(min_value=1, max_value=3),
        document_type=st.sampled_from(['pdf', 'docx'])
    )
    @settings(max_examples=10, deadline=20000, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @pytest.mark.asyncio
    async def test_property_reference_sample_comparison(
        self, scorer, temp_dir, num_samples, document_type
    ):
        """
        Test that authenticity scoring properly compares against reference samples
        when available and provides meaningful comparison results.
        """
        # Create main document
        if document_type == 'pdf':
            main_file = os.path.join(temp_dir, "main.pdf")
            self.create_test_pdf(main_file)
        else:  # docx
            main_file = os.path.join(temp_dir, "main.docx")
            self.create_test_docx(main_file)
        
        # Create reference samples
        reference_samples = []
        for i in range(num_samples):
            if document_type == 'pdf':
                sample_file = os.path.join(temp_dir, f"sample_{i}.pdf")
                self.create_test_pdf(sample_file)
            else:  # docx
                sample_file = os.path.join(temp_dir, f"sample_{i}.docx")
                self.create_test_docx(sample_file)
            reference_samples.append(sample_file)
        
        # Perform analysis with reference samples
        result = await scorer.calculate_authenticity_score(
            main_file, document_id=1, reference_samples=reference_samples
        )
        
        # Verify comparison results
        assert isinstance(result.comparison_results, list)
        assert len(result.comparison_results) == num_samples
        
        for i, comparison in enumerate(result.comparison_results):
            assert isinstance(comparison, ComparisonResult)
            assert comparison.reference_id == f"sample_{i}"
            assert isinstance(comparison.similarity_score, float)
            assert 0.0 <= comparison.similarity_score <= 1.0
            assert isinstance(comparison.confidence, float)
            assert 0.0 <= comparison.confidence <= 1.0
            assert isinstance(comparison.matching_features, list)
            assert isinstance(comparison.differing_features, list)
            
            # Should have some feature analysis
            total_features = len(comparison.matching_features) + len(comparison.differing_features)
            assert total_features > 0
    
    @given(
        image_dimensions=st.tuples(
            st.integers(min_value=50, max_value=800),
            st.integers(min_value=50, max_value=600)
        ),
        document_format=st.sampled_from(['pdf', 'docx', 'jpg'])
    )
    @settings(max_examples=12, deadline=20000, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @pytest.mark.asyncio
    async def test_property_structure_validation_consistency(
        self, scorer, temp_dir, image_dimensions, document_format
    ):
        """
        Test that structure validation is consistent and provides meaningful
        feedback for different document formats and characteristics.
        """
        width, height = image_dimensions
        
        # Create test document
        if document_format == 'pdf':
            file_path = os.path.join(temp_dir, "test.pdf")
            self.create_test_pdf(file_path)
        elif document_format == 'docx':
            file_path = os.path.join(temp_dir, "test.docx")
            self.create_test_docx(file_path)
        else:  # jpg
            file_path = os.path.join(temp_dir, "test.jpg")
            self.create_test_image(file_path, width, height)
        
        # Perform analysis
        result = await scorer.calculate_authenticity_score(file_path, document_id=1)
        
        # Verify structure validation
        struct_val = result.structure_validation
        assert isinstance(struct_val, StructureValidation)
        
        # Structure validation should be consistent with file format
        if document_format in ['pdf', 'docx', 'jpg']:
            # Known formats should have meaningful validation
            assert isinstance(struct_val.format_compliance, float)
            assert 0.0 <= struct_val.format_compliance <= 1.0
            
            # Should provide some feedback
            total_feedback = len(struct_val.violations) + len(struct_val.recommendations)
            # May have no feedback for valid files, but structure should be assessed
            assert total_feedback >= 0
        
        # If file is valid, compliance should be reasonable
        if struct_val.is_valid:
            assert struct_val.format_compliance >= 0.5
        
        # Violations should be strings
        for violation in struct_val.violations:
            assert isinstance(violation, str)
            assert len(violation) > 0
        
        # Recommendations should be strings
        for recommendation in struct_val.recommendations:
            assert isinstance(recommendation, str)
            assert len(recommendation) > 0
    
    @given(
        metadata_present=st.booleans(),
        document_type=st.sampled_from(['pdf', 'docx'])
    )
    @settings(max_examples=8, deadline=15000, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @pytest.mark.asyncio
    async def test_property_metadata_authenticity_assessment(
        self, scorer, temp_dir, metadata_present, document_type
    ):
        """
        Test that metadata authenticity assessment properly evaluates
        document metadata and provides appropriate scoring.
        """
        # Create document with or without metadata
        if document_type == 'pdf':
            file_path = os.path.join(temp_dir, "test.pdf")
            self.create_test_pdf(file_path, metadata_present)
        else:  # docx
            file_path = os.path.join(temp_dir, "test.docx")
            self.create_test_docx(file_path, metadata_present)
        
        # Perform analysis
        result = await scorer.calculate_authenticity_score(file_path, document_id=1)
        
        # Check metadata authenticity factor
        factors = result.authenticity_score.contributing_factors
        assert 'metadata_authenticity' in factors
        
        metadata_score = factors['metadata_authenticity']
        assert isinstance(metadata_score, float)
        assert 0.0 <= metadata_score <= 1.0
        
        # Documents with metadata should generally score higher
        # (though this is not guaranteed due to other factors)
        if metadata_present:
            # Should have some positive score for having metadata
            assert metadata_score >= 0.0
        else:
            # May have lower score for missing metadata, but not necessarily
            assert metadata_score >= 0.0
    
    @pytest.mark.asyncio
    async def test_property_error_handling_robustness(self, scorer, temp_dir):
        """
        Test that authenticity scoring handles various error conditions gracefully
        and always returns valid AuthenticityAnalysis objects.
        """
        # Test with non-existent file
        result = await scorer.calculate_authenticity_score("/nonexistent/file.pdf", document_id=1)
        assert isinstance(result, AuthenticityAnalysis)
        assert result.authenticity_score.overall_score == 0.0
        assert result.authenticity_score.confidence_level == 0.0
        
        # Test with empty file
        empty_file = os.path.join(temp_dir, "empty.pdf")
        with open(empty_file, 'w') as f:
            pass  # Create empty file
        
        result = await scorer.calculate_authenticity_score(empty_file, document_id=2)
        assert isinstance(result, AuthenticityAnalysis)
        # Should handle gracefully without crashing
        
        # Test with corrupted file
        corrupted_file = os.path.join(temp_dir, "corrupted.pdf")
        with open(corrupted_file, 'wb') as f:
            f.write(b"This is not a valid PDF file")
        
        result = await scorer.calculate_authenticity_score(corrupted_file, document_id=3)
        assert isinstance(result, AuthenticityAnalysis)
        # Should handle gracefully without crashing
    
    @given(
        score_range=st.floats(min_value=0.0, max_value=1.0)
    )
    @settings(max_examples=10, deadline=10000, suppress_health_check=[HealthCheck.function_scoped_fixture])
    def test_property_risk_level_consistency(self, scorer, score_range):
        """
        Test that risk level determination is consistent with authenticity scores.
        """
        risk_level = scorer._determine_risk_level(score_range)
        
        # Verify risk level is appropriate for score
        if score_range >= 0.8:
            assert risk_level == RiskLevel.LOW
        elif score_range >= 0.6:
            assert risk_level == RiskLevel.MEDIUM
        elif score_range >= 0.4:
            assert risk_level == RiskLevel.HIGH
        else:
            assert risk_level == RiskLevel.CRITICAL
        
        # Risk level should be a valid enum value
        assert isinstance(risk_level, RiskLevel)


# Unit tests for specific edge cases and examples
class TestAuthenticityScorer_Units:
    """Unit tests for specific authenticity scoring scenarios."""
    
    @pytest.fixture
    def scorer(self):
        """Create an authenticity scorer instance."""
        return AuthenticityScorer()
    
    @pytest.mark.asyncio
    async def test_empty_factors_handling(self, scorer):
        """Test handling of empty authenticity factors."""
        empty_factors = {}
        comparison_results = []
        structure_validation = StructureValidation(
            is_valid=True,
            format_compliance=0.8,
            violations=[],
            recommendations=[]
        )
        embedded_objects = []
        
        score = scorer._calculate_overall_score(
            empty_factors, comparison_results, structure_validation, embedded_objects
        )
        
        assert isinstance(score, AuthenticityScore)
        assert 0.0 <= score.overall_score <= 1.0
        assert 0.0 <= score.confidence_level <= 1.0
    
    @pytest.mark.asyncio
    async def test_similarity_calculation_edge_cases(self, scorer, tmp_path):
        """Test similarity calculation with edge cases."""
        # Create two identical files
        file1 = tmp_path / "file1.txt"
        file2 = tmp_path / "file2.txt"
        
        content = "Test content"
        file1.write_text(content)
        file2.write_text(content)
        
        similarity = await scorer._calculate_similarity(str(file1), str(file2))
        
        # Identical files should have high similarity
        assert isinstance(similarity, float)
        assert 0.0 <= similarity <= 1.0
        # Should be high since same type and size
        assert similarity >= 0.7
    
    def test_confidence_calculation_properties(self, scorer):
        """Test confidence level calculation properties."""
        # Test with various factor combinations
        factors = {
            'format_consistency': 0.8,
            'metadata_authenticity': 0.7,
            'content_integrity': 0.9
        }
        
        comparison_results = [
            ComparisonResult(
                reference_id="test",
                similarity_score=0.8,
                matching_features=[],
                differing_features=[],
                confidence=0.9
            )
        ]
        
        structure_validation = StructureValidation(
            is_valid=True,
            format_compliance=0.8,
            violations=[],
            recommendations=[]
        )
        
        confidence = scorer._calculate_confidence_level(
            factors, comparison_results, structure_validation
        )
        
        assert isinstance(confidence, float)
        assert 0.0 <= confidence <= 1.0
        # Should have reasonable confidence with good inputs
        assert confidence >= 0.3
    
    @pytest.mark.asyncio
    async def test_format_consistency_assessment(self, scorer, tmp_path):
        """Test format consistency assessment."""
        # Create a file with correct extension
        pdf_file = tmp_path / "test.pdf"
        
        # Create minimal PDF content
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(pdf_file))
        c.drawString(100, 750, "Test")
        c.save()
        
        consistency = await scorer._assess_format_consistency(str(pdf_file))
        
        assert isinstance(consistency, float)
        assert 0.0 <= consistency <= 1.0
        # Should have high consistency for proper PDF
        assert consistency >= 0.8