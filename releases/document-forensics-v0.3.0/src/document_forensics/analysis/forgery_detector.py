"""Enhanced forgery detection for multiple document formats.

This module provides specialized forgery detection capabilities for:
- Word documents (.docx, .doc)
- Excel spreadsheets (.xlsx, .xls)
- Text files (.txt)
- Images (.jpg, .png, .tiff, .bmp)
- PDF documents (.pdf)
"""

import io
import logging
import re
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from pathlib import Path

import numpy as np
import cv2
from PIL import Image
import PyPDF2
from docx import Document as DocxDocument
from openpyxl import load_workbook
import xml.etree.ElementTree as ET

from ..core.models import (
    ForgeryAnalysis, ForgeryIndicator, RiskLevel,
    DocumentType, ForgeryType
)

logger = logging.getLogger(__name__)


class ForgeryDetector:
    """Enhanced forgery detection for multiple document formats."""
    
    def __init__(self):
        """Initialize the forgery detector."""
        self.format_detectors = {
            'word': self._detect_word_forgery,
            'excel': self._detect_excel_forgery,
            'text': self._detect_text_forgery,
            'image': self._detect_image_forgery,
            'pdf': self._detect_pdf_forgery
        }
    
    async def detect_forgery(
        self, 
        document_path: str,
        document_id: int
    ) -> ForgeryAnalysis:
        """
        Detect forgery in a document based on its format.
        
        Args:
            document_path: Path to the document
            document_id: Document ID for tracking
            
        Returns:
            ForgeryAnalysis with detected forgery indicators
        """
        logger.info(f"Starting forgery detection for document {document_id}")
        
        try:
            # Determine document type
            file_extension = Path(document_path).suffix.lower()
            doc_type = self._determine_document_type(file_extension)
            
            # Select appropriate detector
            detector = self.format_detectors.get(doc_type)
            if not detector:
                logger.warning(f"No forgery detector for type: {doc_type}")
                return ForgeryAnalysis(
                    document_id=document_id,
                    document_type=doc_type,
                    overall_risk=RiskLevel.LOW,
                    confidence_score=0.0,
                    indicators=[]
                )
            
            # Run format-specific detection
            indicators = await detector(document_path)
            
            # Calculate overall risk
            overall_risk = self._calculate_overall_risk(indicators)
            
            # Calculate confidence score
            confidence_score = self._calculate_confidence(indicators)
            
            return ForgeryAnalysis(
                document_id=document_id,
                document_type=doc_type,
                overall_risk=overall_risk,
                confidence_score=confidence_score,
                indicators=indicators,
                detection_methods_used=self._get_methods_used(indicators)
            )
            
        except Exception as e:
            logger.error(f"Error in forgery detection: {str(e)}")
            return ForgeryAnalysis(
                document_id=document_id,
                document_type='unknown',
                overall_risk=RiskLevel.LOW,
                confidence_score=0.0,
                indicators=[],
                error_message=str(e)
            )

    def _determine_document_type(self, file_extension: str) -> str:
        """Determine document type from file extension."""
        extension_map = {
            '.docx': 'word', '.doc': 'word',
            '.xlsx': 'excel', '.xls': 'excel',
            '.txt': 'text', '.text': 'text',
            '.jpg': 'image', '.jpeg': 'image', '.png': 'image',
            '.tiff': 'image', '.tif': 'image', '.bmp': 'image',
            '.pdf': 'pdf'
        }
        return extension_map.get(file_extension, 'unknown')
    
    async def _detect_word_forgery(self, document_path: str) -> List[ForgeryIndicator]:
        """Detect forgery indicators in Word documents."""
        indicators = []
        
        try:
            doc = DocxDocument(document_path)
            
            # Check 1: Revision history analysis
            revision_indicators = await self._analyze_word_revisions(doc)
            indicators.extend(revision_indicators)
            
            # Check 2: Style inconsistencies
            style_indicators = await self._analyze_word_styles(doc)
            indicators.extend(style_indicators)
            
            # Check 3: Font manipulation detection
            font_indicators = await self._analyze_word_fonts(doc)
            indicators.extend(font_indicators)
            
            # Check 4: Hidden text detection
            hidden_indicators = await self._detect_hidden_text_word(doc)
            indicators.extend(hidden_indicators)
            
            # Check 5: Track changes analysis
            track_changes_indicators = await self._analyze_track_changes(doc)
            indicators.extend(track_changes_indicators)
            
            # Check 6: XML structure analysis
            xml_indicators = await self._analyze_word_xml(document_path)
            indicators.extend(xml_indicators)
            
        except Exception as e:
            logger.error(f"Error in Word forgery detection: {str(e)}")
            indicators.append(ForgeryIndicator(
                type=ForgeryType.ANALYSIS_ERROR,
                description=f"Error analyzing Word document: {str(e)}",
                confidence=0.5,
                severity=RiskLevel.MEDIUM
            ))
        
        return indicators

    async def _detect_excel_forgery(self, document_path: str) -> List[ForgeryIndicator]:
        """Detect forgery indicators in Excel spreadsheets."""
        indicators = []
        
        try:
            workbook = load_workbook(document_path)
            
            # Check 1: Formula manipulation detection
            formula_indicators = await self._analyze_excel_formulas(workbook)
            indicators.extend(formula_indicators)
            
            # Check 2: Cell value inconsistencies
            value_indicators = await self._analyze_cell_values(workbook)
            indicators.extend(value_indicators)
            
            # Check 3: Hidden sheets/rows/columns
            hidden_indicators = await self._detect_hidden_content_excel(workbook)
            indicators.extend(hidden_indicators)
            
            # Check 4: Data validation tampering
            validation_indicators = await self._analyze_data_validation(workbook)
            indicators.extend(validation_indicators)
            
            # Check 5: Macro analysis
            macro_indicators = await self._analyze_excel_macros(workbook)
            indicators.extend(macro_indicators)
            
            # Check 6: Number format manipulation
            format_indicators = await self._analyze_number_formats(workbook)
            indicators.extend(format_indicators)
            
        except Exception as e:
            logger.error(f"Error in Excel forgery detection: {str(e)}")
            indicators.append(ForgeryIndicator(
                type=ForgeryType.ANALYSIS_ERROR,
                description=f"Error analyzing Excel document: {str(e)}",
                confidence=0.5,
                severity=RiskLevel.MEDIUM
            ))
        
        return indicators

    async def _detect_text_forgery(self, document_path: str) -> List[ForgeryIndicator]:
        """Detect forgery indicators in text files."""
        indicators = []
        
        try:
            with open(document_path, 'rb') as f:
                content = f.read()
            
            # Check 1: Encoding manipulation
            encoding_indicators = await self._analyze_text_encoding(content)
            indicators.extend(encoding_indicators)
            
            # Check 2: Invisible characters
            invisible_indicators = await self._detect_invisible_characters(content)
            indicators.extend(invisible_indicators)
            
            # Check 3: Line ending inconsistencies
            line_ending_indicators = await self._analyze_line_endings(content)
            indicators.extend(line_ending_indicators)
            
            # Check 4: Homoglyph attacks
            homoglyph_indicators = await self._detect_homoglyphs(content)
            indicators.extend(homoglyph_indicators)
            
        except Exception as e:
            logger.error(f"Error in text forgery detection: {str(e)}")
        
        return indicators
    
    async def _detect_image_forgery(self, document_path: str) -> List[ForgeryIndicator]:
        """Detect forgery indicators in images."""
        indicators = []
        
        try:
            image = cv2.imread(document_path)
            if image is None:
                return indicators
            
            # Check 1: Clone detection
            clone_indicators = await self._detect_cloning(image)
            indicators.extend(clone_indicators)
            
            # Check 2: Noise analysis
            noise_indicators = await self._analyze_image_noise(image)
            indicators.extend(noise_indicators)
            
            # Check 3: Compression artifacts
            compression_indicators = await self._analyze_compression(document_path)
            indicators.extend(compression_indicators)
            
            # Check 4: Lighting inconsistencies
            lighting_indicators = await self._analyze_lighting(image)
            indicators.extend(lighting_indicators)
            
            # Check 5: Edge analysis
            edge_indicators = await self._analyze_edges(image)
            indicators.extend(edge_indicators)
            
        except Exception as e:
            logger.error(f"Error in image forgery detection: {str(e)}")
        
        return indicators
    
    async def _detect_pdf_forgery(self, document_path: str) -> List[ForgeryIndicator]:
        """Detect forgery indicators in PDF documents."""
        indicators = []
        
        try:
            with open(document_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Check 1: Digital signature verification
                signature_indicators = await self._verify_pdf_signatures(pdf_reader)
                indicators.extend(signature_indicators)
                
                # Check 2: Incremental updates
                update_indicators = await self._analyze_incremental_updates(document_path)
                indicators.extend(update_indicators)
                
                # Check 3: Object stream analysis
                object_indicators = await self._analyze_pdf_objects(pdf_reader)
                indicators.extend(object_indicators)
                
                # Check 4: Text layer analysis
                text_indicators = await self._analyze_pdf_text_layer(pdf_reader)
                indicators.extend(text_indicators)
                
                # Check 5: Form field tampering
                form_indicators = await self._analyze_pdf_forms(pdf_reader)
                indicators.extend(form_indicators)
                
        except Exception as e:
            logger.error(f"Error in PDF forgery detection: {str(e)}")
        
        return indicators

    # Word Document Detection Helper Methods
    async def _analyze_word_revisions(self, doc: DocxDocument) -> List['ForgeryIndicator']:
        """Analyze revision history for suspicious patterns."""
        indicators = []
        
        try:
            # Access core properties
            props = doc.core_properties
            
            # Check for multiple authors
            if props.last_modified_by and props.author:
                if props.last_modified_by != props.author:
                    indicators.append(self._create_indicator(
                        type='REVISION_MANIPULATION',
                        description=f"Document modified by different user: {props.last_modified_by} (original: {props.author})",
                        confidence=0.6,
                        severity='MEDIUM',
                        evidence={'author': props.author, 'last_modified_by': props.last_modified_by}
                    ))
            
            # Check revision number
            if props.revision and int(props.revision) > 100:
                indicators.append(self._create_indicator(
                    type='REVISION_MANIPULATION',
                    description=f"Unusually high revision count: {props.revision}",
                    confidence=0.5,
                    severity='LOW',
                    evidence={'revision_count': props.revision}
                ))
                
        except Exception as e:
            logger.error(f"Error analyzing Word revisions: {str(e)}")
        
        return indicators
    
    async def _analyze_word_styles(self, doc: DocxDocument) -> List['ForgeryIndicator']:
        """Detect style inconsistencies."""
        indicators = []
        
        try:
            style_counts = {}
            
            for paragraph in doc.paragraphs:
                if paragraph.style:
                    style_name = paragraph.style.name
                    style_counts[style_name] = style_counts.get(style_name, 0) + 1
            
            # Check for excessive style variation
            if len(style_counts) > 10:
                indicators.append(self._create_indicator(
                    type='STYLE_INCONSISTENCY',
                    description=f"Excessive style variation detected: {len(style_counts)} different styles",
                    confidence=0.6,
                    severity='MEDIUM',
                    evidence={'style_count': len(style_counts), 'styles': list(style_counts.keys())}
                ))
                
        except Exception as e:
            logger.error(f"Error analyzing Word styles: {str(e)}")
        
        return indicators
    
    async def _analyze_word_fonts(self, doc: DocxDocument) -> List['ForgeryIndicator']:
        """Detect font manipulation."""
        indicators = []
        
        try:
            for para_idx, paragraph in enumerate(doc.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    # Check for white text (potential hidden content)
                    if run.font.color and hasattr(run.font.color, 'rgb'):
                        if run.font.color.rgb == (255, 255, 255) or run.font.color.rgb == 'FFFFFF':
                            indicators.append(self._create_indicator(
                                type='FONT_MANIPULATION',
                                description=f"White text detected in paragraph {para_idx + 1}",
                                confidence=0.9,
                                severity='HIGH',
                                location={'paragraph': para_idx + 1, 'run': run_idx + 1},
                                evidence={'text_content': run.text[:50]}
                            ))
                    
                    # Check for very small font sizes
                    if run.font.size and run.font.size.pt < 4:
                        indicators.append(self._create_indicator(
                            type='FONT_MANIPULATION',
                            description=f"Extremely small font size detected: {run.font.size.pt}pt",
                            confidence=0.7,
                            severity='MEDIUM',
                            location={'paragraph': para_idx + 1, 'run': run_idx + 1}
                        ))
                        
        except Exception as e:
            logger.error(f"Error analyzing Word fonts: {str(e)}")
        
        return indicators
    
    async def _detect_hidden_text_word(self, doc: DocxDocument) -> List['ForgeryIndicator']:
        """Find hidden text."""
        indicators = []
        
        try:
            for para_idx, paragraph in enumerate(doc.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    # Check if run is marked as hidden
                    if run.font.hidden:
                        indicators.append(self._create_indicator(
                            type='HIDDEN_TEXT',
                            description=f"Hidden text found in paragraph {para_idx + 1}",
                            confidence=0.95,
                            severity='HIGH',
                            location={'paragraph': para_idx + 1, 'run': run_idx + 1},
                            evidence={'text_content': run.text[:100]}
                        ))
                        
        except Exception as e:
            logger.error(f"Error detecting hidden text: {str(e)}")
        
        return indicators
    
    async def _analyze_track_changes(self, doc: DocxDocument) -> List['ForgeryIndicator']:
        """Analyze track changes."""
        indicators = []
        
        try:
            # Check if track changes are enabled
            # Note: python-docx has limited track changes support
            # This is a simplified implementation
            
            # Check for deleted content in XML
            from docx.oxml import parse_xml
            for paragraph in doc.paragraphs:
                if paragraph._element.xml:
                    xml_str = paragraph._element.xml.decode() if isinstance(paragraph._element.xml, bytes) else str(paragraph._element.xml)
                    if 'w:del' in xml_str or 'w:ins' in xml_str:
                        indicators.append(self._create_indicator(
                            type='TRACK_CHANGES_ANOMALY',
                            description="Track changes detected in document",
                            confidence=0.8,
                            severity='MEDIUM',
                            evidence={'has_track_changes': True}
                        ))
                        break
                        
        except Exception as e:
            logger.error(f"Error analyzing track changes: {str(e)}")
        
        return indicators
    
    async def _analyze_word_xml(self, document_path: str) -> List['ForgeryIndicator']:
        """Analyze XML structure."""
        indicators = []
        
        try:
            import zipfile
            
            with zipfile.ZipFile(document_path, 'r') as docx_zip:
                # Check for suspicious files in the package
                file_list = docx_zip.namelist()
                
                # Look for unusual files
                suspicious_files = [f for f in file_list if f.endswith('.exe') or f.endswith('.dll') or f.endswith('.vbs')]
                if suspicious_files:
                    indicators.append(self._create_indicator(
                        type='XML_STRUCTURE_ANOMALY',
                        description=f"Suspicious files found in document package: {suspicious_files}",
                        confidence=0.95,
                        severity='CRITICAL',
                        evidence={'suspicious_files': suspicious_files}
                    ))
                    
        except Exception as e:
            logger.error(f"Error analyzing Word XML: {str(e)}")
        
        return indicators

    # Excel Detection Helper Methods
    async def _analyze_excel_formulas(self, workbook) -> List['ForgeryIndicator']:
        """Detect formula tampering."""
        indicators = []
        
        try:
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        # Check if cell has a formula
                        if cell.data_type == 'f':
                            # Check if formula result seems inconsistent
                            if cell.value is not None:
                                # This is a simplified check
                                formula_str = str(cell.value)
                                if 'ERROR' in formula_str.upper() or '#' in formula_str:
                                    indicators.append(self._create_indicator(
                                        type='FORMULA_TAMPERING',
                                        description=f"Formula error in cell {cell.coordinate}",
                                        confidence=0.7,
                                        severity='MEDIUM',
                                        location={'sheet': sheet.title, 'cell': cell.coordinate},
                                        evidence={'formula': formula_str}
                                    ))
                                    
        except Exception as e:
            logger.error(f"Error analyzing Excel formulas: {str(e)}")
        
        return indicators
    
    async def _analyze_cell_values(self, workbook) -> List['ForgeryIndicator']:
        """Find value inconsistencies."""
        indicators = []
        
        try:
            for sheet in workbook.worksheets:
                # Check for cells with formulas that have been overridden with values
                for row in sheet.iter_rows():
                    for cell in row:
                        # Check if cell appears to be a calculated value but has no formula
                        if cell.data_type == 'n' and cell.value is not None:
                            # Look for patterns that suggest manual override
                            # This is a heuristic check
                            if isinstance(cell.value, (int, float)):
                                # Check neighboring cells for formulas
                                neighbors_have_formulas = False
                                try:
                                    for offset in [(-1, 0), (1, 0), (0, -1), (0, 1)]:
                                        neighbor = sheet.cell(cell.row + offset[0], cell.column + offset[1])
                                        if neighbor.data_type == 'f':
                                            neighbors_have_formulas = True
                                            break
                                except:
                                    pass
                                
                                if neighbors_have_formulas:
                                    indicators.append(self._create_indicator(
                                        type='VALUE_OVERRIDE',
                                        description=f"Potential manual value override in cell {cell.coordinate}",
                                        confidence=0.5,
                                        severity='LOW',
                                        location={'sheet': sheet.title, 'cell': cell.coordinate}
                                    ))
                                    
        except Exception as e:
            logger.error(f"Error analyzing cell values: {str(e)}")
        
        return indicators
    
    async def _detect_hidden_content_excel(self, workbook) -> List['ForgeryIndicator']:
        """Detect hidden sheets/rows/columns."""
        indicators = []
        
        try:
            # Check for hidden sheets
            for sheet in workbook.worksheets:
                if sheet.sheet_state == 'hidden':
                    indicators.append(self._create_indicator(
                        type='HIDDEN_CONTENT',
                        description=f"Hidden sheet detected: {sheet.title}",
                        confidence=0.9,
                        severity='HIGH',
                        location={'sheet': sheet.title},
                        evidence={'sheet_state': sheet.sheet_state}
                    ))
                
                # Check for hidden rows/columns
                hidden_rows = sum(1 for row in sheet.row_dimensions.values() if row.hidden)
                hidden_cols = sum(1 for col in sheet.column_dimensions.values() if col.hidden)
                
                if hidden_rows > 0:
                    indicators.append(self._create_indicator(
                        type='HIDDEN_CONTENT',
                        description=f"Hidden rows detected in sheet {sheet.title}: {hidden_rows} rows",
                        confidence=0.8,
                        severity='MEDIUM',
                        location={'sheet': sheet.title},
                        evidence={'hidden_rows': hidden_rows}
                    ))
                
                if hidden_cols > 0:
                    indicators.append(self._create_indicator(
                        type='HIDDEN_CONTENT',
                        description=f"Hidden columns detected in sheet {sheet.title}: {hidden_cols} columns",
                        confidence=0.8,
                        severity='MEDIUM',
                        location={'sheet': sheet.title},
                        evidence={'hidden_columns': hidden_cols}
                    ))
                    
        except Exception as e:
            logger.error(f"Error detecting hidden content: {str(e)}")
        
        return indicators
    
    async def _analyze_data_validation(self, workbook) -> List['ForgeryIndicator']:
        """Check data validation tampering."""
        indicators = []
        
        try:
            for sheet in workbook.worksheets:
                if sheet.data_validations:
                    for validation in sheet.data_validations.dataValidation:
                        # Check if validation is disabled
                        if validation.showErrorMessage == False:
                            indicators.append(self._create_indicator(
                                type='VALIDATION_BYPASS',
                                description=f"Data validation disabled in sheet {sheet.title}",
                                confidence=0.7,
                                severity='MEDIUM',
                                location={'sheet': sheet.title},
                                evidence={'validation_disabled': True}
                            ))
                            
        except Exception as e:
            logger.error(f"Error analyzing data validation: {str(e)}")
        
        return indicators
    
    async def _analyze_excel_macros(self, workbook) -> List['ForgeryIndicator']:
        """Analyze VBA macros."""
        indicators = []
        
        try:
            # Check if workbook has VBA project
            if hasattr(workbook, 'vba_archive') and workbook.vba_archive:
                indicators.append(self._create_indicator(
                    type='MACRO_SUSPICIOUS',
                    description="VBA macros detected in workbook",
                    confidence=0.6,
                    severity='MEDIUM',
                    evidence={'has_macros': True}
                ))
                
        except Exception as e:
            logger.error(f"Error analyzing macros: {str(e)}")
        
        return indicators
    
    async def _analyze_number_formats(self, workbook) -> List['ForgeryIndicator']:
        """Check number format manipulation."""
        indicators = []
        
        try:
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.number_format and cell.number_format != 'General':
                            # Check for suspicious custom formats
                            if ';;;' in cell.number_format or '[<0]' in cell.number_format:
                                indicators.append(self._create_indicator(
                                    type='NUMBER_FORMAT_MANIPULATION',
                                    description=f"Suspicious number format in cell {cell.coordinate}",
                                    confidence=0.6,
                                    severity='LOW',
                                    location={'sheet': sheet.title, 'cell': cell.coordinate},
                                    evidence={'number_format': cell.number_format}
                                ))
                                
        except Exception as e:
            logger.error(f"Error analyzing number formats: {str(e)}")
        
        return indicators

    # Text File Detection Helper Methods
    async def _analyze_text_encoding(self, content: bytes) -> List['ForgeryIndicator']:
        """Detect encoding manipulation."""
        indicators = []
        
        try:
            import chardet
            
            # Detect encoding
            detected = chardet.detect(content)
            encoding = detected.get('encoding', 'unknown')
            confidence = detected.get('confidence', 0.0)
            
            # Check for mixed encodings (simplified check)
            try:
                content.decode('utf-8')
                utf8_valid = True
            except:
                utf8_valid = False
            
            try:
                content.decode('ascii')
                ascii_valid = True
            except:
                ascii_valid = False
            
            if not utf8_valid and not ascii_valid and encoding not in ['utf-8', 'ascii']:
                indicators.append(self._create_indicator(
                    type='ENCODING_MANIPULATION',
                    description=f"Unusual encoding detected: {encoding}",
                    confidence=0.7,
                    severity='MEDIUM',
                    evidence={'encoding': encoding, 'detection_confidence': confidence}
                ))
                
        except Exception as e:
            logger.error(f"Error analyzing text encoding: {str(e)}")
        
        return indicators
    
    async def _detect_invisible_characters(self, content: bytes) -> List['ForgeryIndicator']:
        """Find zero-width and control characters."""
        indicators = []
        
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Zero-width characters
            zero_width_chars = [
                '\u200B',  # Zero-width space
                '\u200C',  # Zero-width non-joiner
                '\u200D',  # Zero-width joiner
                '\uFEFF',  # Zero-width no-break space
            ]
            
            for char in zero_width_chars:
                if char in text:
                    count = text.count(char)
                    indicators.append(self._create_indicator(
                        type='INVISIBLE_CHARACTERS',
                        description=f"Zero-width character detected: U+{ord(char):04X} ({count} occurrences)",
                        confidence=0.9,
                        severity='HIGH',
                        evidence={'character_code': f"U+{ord(char):04X}", 'count': count}
                    ))
                    
        except Exception as e:
            logger.error(f"Error detecting invisible characters: {str(e)}")
        
        return indicators
    
    async def _analyze_line_endings(self, content: bytes) -> List['ForgeryIndicator']:
        """Check line ending consistency."""
        indicators = []
        
        try:
            crlf_count = content.count(b'\r\n')
            lf_count = content.count(b'\n') - crlf_count
            cr_count = content.count(b'\r') - crlf_count
            
            # Check for mixed line endings
            endings_present = sum([crlf_count > 0, lf_count > 0, cr_count > 0])
            
            if endings_present > 1:
                indicators.append(self._create_indicator(
                    type='LINE_ENDING_INCONSISTENCY',
                    description=f"Mixed line endings detected (CRLF: {crlf_count}, LF: {lf_count}, CR: {cr_count})",
                    confidence=0.7,
                    severity='LOW',
                    evidence={'crlf': crlf_count, 'lf': lf_count, 'cr': cr_count}
                ))
                
        except Exception as e:
            logger.error(f"Error analyzing line endings: {str(e)}")
        
        return indicators
    
    async def _detect_homoglyphs(self, content: bytes) -> List['ForgeryIndicator']:
        """Detect look-alike character substitutions."""
        indicators = []
        
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Common homoglyph pairs (Latin vs Cyrillic)
            homoglyphs = {
                'а': 'a',  # Cyrillic 'a' vs Latin 'a'
                'е': 'e',  # Cyrillic 'e' vs Latin 'e'
                'о': 'o',  # Cyrillic 'o' vs Latin 'o'
                'р': 'p',  # Cyrillic 'r' vs Latin 'p'
                'с': 'c',  # Cyrillic 's' vs Latin 'c'
                'у': 'y',  # Cyrillic 'u' vs Latin 'y'
                'х': 'x',  # Cyrillic 'h' vs Latin 'x'
            }
            
            for cyrillic, latin in homoglyphs.items():
                if cyrillic in text:
                    count = text.count(cyrillic)
                    indicators.append(self._create_indicator(
                        type='HOMOGLYPH_ATTACK',
                        description=f"Cyrillic character '{cyrillic}' (looks like '{latin}') detected ({count} occurrences)",
                        confidence=0.8,
                        severity='MEDIUM',
                        evidence={'cyrillic_char': cyrillic, 'latin_lookalike': latin, 'count': count}
                    ))
                    
        except Exception as e:
            logger.error(f"Error detecting homoglyphs: {str(e)}")
        
        return indicators

    # Image Detection Helper Methods
    async def _detect_cloning(self, image: np.ndarray) -> List['ForgeryIndicator']:
        """Detect cloned regions."""
        indicators = []
        
        try:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
            
            # Use ORB feature detection to find similar regions
            orb = cv2.ORB_create(nfeatures=1000)
            keypoints, descriptors = orb.detectAndCompute(gray, None)
            
            if descriptors is not None and len(descriptors) > 10:
                # Use BFMatcher to find similar features
                bf = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=False)
                matches = bf.knnMatch(descriptors, descriptors, k=3)
                
                # Count self-matches (excluding identity matches)
                similar_regions = 0
                for match_group in matches:
                    if len(match_group) >= 2:
                        # Check if there are very similar features at different locations
                        if match_group[1].distance < 50:  # Threshold for similarity
                            similar_regions += 1
                
                if similar_regions > 50:  # Threshold for suspicious cloning
                    indicators.append(self._create_indicator(
                        type='CLONE_DETECTION',
                        description=f"Potential cloned regions detected ({similar_regions} similar feature groups)",
                        confidence=0.7,
                        severity='MEDIUM',
                        evidence={'similar_regions': similar_regions}
                    ))
                    
        except Exception as e:
            logger.error(f"Error detecting cloning: {str(e)}")
        
        return indicators
    
    async def _analyze_image_noise(self, image: np.ndarray) -> List['ForgeryIndicator']:
        """Analyze noise patterns."""
        indicators = []
        
        try:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
            
            # Calculate noise variance in different regions
            h, w = gray.shape
            block_size = 64
            variances = []
            
            for y in range(0, h - block_size, block_size):
                for x in range(0, w - block_size, block_size):
                    block = gray[y:y+block_size, x:x+block_size]
                    variances.append(np.var(block))
            
            if variances:
                mean_var = np.mean(variances)
                std_var = np.std(variances)
                
                # Check for unusual variance distribution
                if std_var > mean_var * 0.5:  # High variation in noise
                    indicators.append(self._create_indicator(
                        type='NOISE_INCONSISTENCY',
                        description="Inconsistent noise patterns detected across image regions",
                        confidence=0.6,
                        severity='MEDIUM',
                        evidence={'mean_variance': float(mean_var), 'std_variance': float(std_var)}
                    ))
                    
        except Exception as e:
            logger.error(f"Error analyzing image noise: {str(e)}")
        
        return indicators
    
    async def _analyze_compression(self, image_path: str) -> List['ForgeryIndicator']:
        """Detect compression artifacts."""
        indicators = []
        
        try:
            if not image_path.lower().endswith(('.jpg', '.jpeg')):
                return indicators
            
            image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
            if image is None:
                return indicators
            
            # Analyze DCT coefficients for double compression
            h, w = image.shape
            block_size = 8
            dct_values = []
            
            for y in range(0, h - block_size, block_size):
                for x in range(0, w - block_size, block_size):
                    block = image[y:y+block_size, x:x+block_size].astype(float)
                    dct_block = cv2.dct(block)
                    dct_values.append(np.abs(dct_block).flatten())
            
            if dct_values:
                # Check for double quantization artifacts
                all_dct = np.concatenate(dct_values)
                hist, _ = np.histogram(all_dct, bins=50)
                
                # Look for periodic patterns in histogram (sign of double compression)
                if len(hist) > 10:
                    peaks = np.where(hist > np.mean(hist) * 2)[0]
                    if len(peaks) > 5:
                        indicators.append(self._create_indicator(
                            type='COMPRESSION_ANOMALY',
                            description="Double JPEG compression detected",
                            confidence=0.7,
                            severity='MEDIUM',
                            evidence={'histogram_peaks': len(peaks)}
                        ))
                        
        except Exception as e:
            logger.error(f"Error analyzing compression: {str(e)}")
        
        return indicators
    
    async def _analyze_lighting(self, image: np.ndarray) -> List['ForgeryIndicator']:
        """Check lighting consistency."""
        indicators = []
        
        try:
            # Convert to LAB color space for better lighting analysis
            lab = cv2.cvtColor(image, cv2.COLOR_BGR2LAB)
            l_channel = lab[:, :, 0]
            
            # Divide image into regions and analyze lighting
            h, w = l_channel.shape
            region_size = min(h, w) // 4
            
            region_means = []
            for y in range(0, h - region_size, region_size):
                for x in range(0, w - region_size, region_size):
                    region = l_channel[y:y+region_size, x:x+region_size]
                    region_means.append(np.mean(region))
            
            if region_means:
                std_lighting = np.std(region_means)
                mean_lighting = np.mean(region_means)
                
                # Check for unusual lighting variation
                if std_lighting > mean_lighting * 0.3:
                    indicators.append(self._create_indicator(
                        type='LIGHTING_INCONSISTENCY',
                        description="Inconsistent lighting detected across image regions",
                        confidence=0.6,
                        severity='MEDIUM',
                        evidence={'lighting_std': float(std_lighting), 'lighting_mean': float(mean_lighting)}
                    ))
                    
        except Exception as e:
            logger.error(f"Error analyzing lighting: {str(e)}")
        
        return indicators
    
    async def _analyze_edges(self, image: np.ndarray) -> List['ForgeryIndicator']:
        """Analyze edge consistency."""
        indicators = []
        
        try:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
            
            # Apply Canny edge detection
            edges = cv2.Canny(gray, 50, 150)
            
            # Analyze edge density in different regions
            h, w = edges.shape
            region_size = min(h, w) // 4
            
            edge_densities = []
            for y in range(0, h - region_size, region_size):
                for x in range(0, w - region_size, region_size):
                    region = edges[y:y+region_size, x:x+region_size]
                    density = np.sum(region > 0) / (region_size * region_size)
                    edge_densities.append(density)
            
            if edge_densities:
                std_density = np.std(edge_densities)
                mean_density = np.mean(edge_densities)
                
                # Check for unusual edge density variation
                if std_density > mean_density * 0.5:
                    indicators.append(self._create_indicator(
                        type='EDGE_INCONSISTENCY',
                        description="Inconsistent edge patterns detected (possible splicing)",
                        confidence=0.6,
                        severity='MEDIUM',
                        evidence={'edge_density_std': float(std_density), 'edge_density_mean': float(mean_density)}
                    ))
                    
        except Exception as e:
            logger.error(f"Error analyzing edges: {str(e)}")
        
        return indicators

    # PDF Detection Helper Methods
    async def _verify_pdf_signatures(self, pdf_reader: PyPDF2.PdfReader) -> List['ForgeryIndicator']:
        """Verify digital signatures."""
        indicators = []
        
        try:
            # Check if PDF has signature fields
            if hasattr(pdf_reader, 'get_fields'):
                fields = pdf_reader.get_fields()
                if fields:
                    for field_name, field_data in fields.items():
                        if field_data.get('/FT') == '/Sig':
                            # Signature field found
                            indicators.append(self._create_indicator(
                                type='SIGNATURE_BROKEN',
                                description=f"Digital signature field detected: {field_name} (verification needed)",
                                confidence=0.5,
                                severity='MEDIUM',
                                evidence={'signature_field': field_name}
                            ))
                            
        except Exception as e:
            logger.error(f"Error verifying PDF signatures: {str(e)}")
        
        return indicators
    
    async def _analyze_incremental_updates(self, document_path: str) -> List['ForgeryIndicator']:
        """Check for post-signature updates."""
        indicators = []
        
        try:
            with open(document_path, 'rb') as f:
                content = f.read()
            
            # Count xref sections (indicates incremental updates)
            xref_count = content.count(b'xref')
            
            if xref_count > 1:
                indicators.append(self._create_indicator(
                    type='INCREMENTAL_UPDATE',
                    description=f"Multiple xref sections detected ({xref_count}), indicating incremental updates",
                    confidence=0.7,
                    severity='MEDIUM',
                    evidence={'xref_count': xref_count}
                ))
                
        except Exception as e:
            logger.error(f"Error analyzing incremental updates: {str(e)}")
        
        return indicators
    
    async def _analyze_pdf_objects(self, pdf_reader: PyPDF2.PdfReader) -> List['ForgeryIndicator']:
        """Analyze PDF object streams."""
        indicators = []
        
        try:
            # Check for suspicious object types
            for page in pdf_reader.pages:
                if '/Contents' in page:
                    # Check for JavaScript
                    page_content = str(page.get('/Contents', ''))
                    if 'JavaScript' in page_content or '/JS' in page_content:
                        indicators.append(self._create_indicator(
                            type='OBJECT_MANIPULATION',
                            description="JavaScript detected in PDF (potential security risk)",
                            confidence=0.8,
                            severity='HIGH',
                            evidence={'has_javascript': True}
                        ))
                        break
                        
        except Exception as e:
            logger.error(f"Error analyzing PDF objects: {str(e)}")
        
        return indicators
    
    async def _analyze_pdf_text_layer(self, pdf_reader: PyPDF2.PdfReader) -> List['ForgeryIndicator']:
        """Compare visible vs extractable text."""
        indicators = []
        
        try:
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                # Check for suspicious patterns
                if text and len(text.strip()) < 10:
                    # Very little extractable text might indicate image-based PDF
                    indicators.append(self._create_indicator(
                        type='TEXT_LAYER_MISMATCH',
                        description=f"Minimal extractable text on page {page_num + 1} (possible scanned/image PDF)",
                        confidence=0.5,
                        severity='LOW',
                        location={'page': page_num + 1},
                        evidence={'text_length': len(text.strip())}
                    ))
                    
        except Exception as e:
            logger.error(f"Error analyzing PDF text layer: {str(e)}")
        
        return indicators
    
    async def _analyze_pdf_forms(self, pdf_reader: PyPDF2.PdfReader) -> List['ForgeryIndicator']:
        """Check form field tampering."""
        indicators = []
        
        try:
            if hasattr(pdf_reader, 'get_fields'):
                fields = pdf_reader.get_fields()
                if fields:
                    # Check for read-only fields that might have been modified
                    for field_name, field_data in fields.items():
                        if field_data.get('/Ff', 0) & 1:  # Read-only flag
                            indicators.append(self._create_indicator(
                                type='FORM_FIELD_TAMPERING',
                                description=f"Read-only form field detected: {field_name}",
                                confidence=0.6,
                                severity='MEDIUM',
                                evidence={'field_name': field_name, 'is_readonly': True}
                            ))
                            
        except Exception as e:
            logger.error(f"Error analyzing PDF forms: {str(e)}")
        
        return indicators
    
    # Utility Methods
    def _calculate_overall_risk(self, indicators: List['ForgeryIndicator']) -> 'RiskLevel':
        """Calculate overall risk level from indicators."""
        if not indicators:
            return 'LOW'
        
        # Count by severity
        critical_count = sum(1 for i in indicators if i['severity'] == 'CRITICAL')
        high_count = sum(1 for i in indicators if i['severity'] == 'HIGH')
        medium_count = sum(1 for i in indicators if i['severity'] == 'MEDIUM')
        
        # Determine overall risk
        if critical_count >= 1:
            return 'CRITICAL'
        elif high_count >= 2:
            return 'CRITICAL'
        elif high_count >= 1:
            return 'HIGH'
        elif medium_count >= 3:
            return 'HIGH'
        elif medium_count >= 1:
            return 'MEDIUM'
        else:
            return 'LOW'
    
    def _calculate_confidence(self, indicators: List['ForgeryIndicator']) -> float:
        """Calculate overall confidence score."""
        if not indicators:
            return 0.0
        
        # Average confidence weighted by severity
        severity_weights = {
            'CRITICAL': 1.0,
            'HIGH': 0.8,
            'MEDIUM': 0.6,
            'LOW': 0.4
        }
        
        weighted_sum = sum(
            i['confidence'] * severity_weights.get(i['severity'], 0.5)
            for i in indicators
        )
        weight_total = sum(
            severity_weights.get(i['severity'], 0.5)
            for i in indicators
        )
        
        return weighted_sum / weight_total if weight_total > 0 else 0.0
    
    def _get_methods_used(self, indicators: List['ForgeryIndicator']) -> List[str]:
        """Get list of detection methods used."""
        methods = set()
        for indicator in indicators:
            if 'detection_method' in indicator:
                methods.add(indicator['detection_method'])
        return list(methods)
    
    def _create_indicator(
        self,
        type: str,
        description: str,
        confidence: float,
        severity: str,
        location: Optional[Dict[str, Any]] = None,
        evidence: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Helper to create forgery indicator."""
        from ..core.models import ForgeryType, RiskLevel
        
        return {
            'type': type,
            'description': description,
            'confidence': confidence,
            'severity': severity,
            'location': location,
            'evidence': evidence or {},
            'detection_method': f"{type.lower()}_detection"
        }
