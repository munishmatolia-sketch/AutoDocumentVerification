"""Metadata extraction and analysis component for document forensics."""

import io
import logging
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import exifread
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

from ..core.models import (
    DeviceFingerprint,
    GeoLocation,
    MetadataAnalysis,
    MetadataAnomaly,
    RiskLevel,
    SoftwareSignature,
    TimestampConsistency,
)

logger = logging.getLogger(__name__)


class MetadataExtractor:
    """Extracts and analyzes metadata from various document types."""

    def __init__(self):
        """Initialize the metadata extractor."""
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif'],
            'pdf': ['.pdf'],
            'docx': ['.docx'],
            'xlsx': ['.xlsx', '.xls']
        }

    async def extract_metadata(self, file_path: str, file_content: bytes) -> MetadataAnalysis:
        """
        Extract comprehensive metadata from a document.
        
        Args:
            file_path: Path to the document file
            file_content: Raw file content as bytes
            
        Returns:
            MetadataAnalysis: Comprehensive metadata analysis results
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            extracted_metadata = {}
            
            # Determine file type and extract appropriate metadata
            if file_extension in self.supported_formats['image']:
                extracted_metadata = await self._extract_image_metadata(file_content)
            elif file_extension in self.supported_formats['pdf']:
                extracted_metadata = await self._extract_pdf_metadata(file_content)
            elif file_extension in self.supported_formats['docx']:
                extracted_metadata = await self._extract_docx_metadata(file_content)
            elif file_extension in self.supported_formats['xlsx']:
                extracted_metadata = await self._extract_xlsx_metadata(file_content)
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                extracted_metadata = {"error": f"Unsupported format: {file_extension}"}

            # Analyze timestamps for consistency
            timestamp_consistency = self._analyze_timestamp_consistency(extracted_metadata)
            
            # Detect software signatures
            software_signatures = self._detect_software_signatures(extracted_metadata)
            
            # Detect anomalies
            anomalies = self._detect_metadata_anomalies(extracted_metadata)
            
            # Extract geo location if available
            geo_location = self._extract_geo_location(extracted_metadata)
            
            # Extract device fingerprint
            device_fingerprint = self._extract_device_fingerprint(extracted_metadata)

            return MetadataAnalysis(
                document_id=0,  # Will be set by caller
                extracted_metadata=extracted_metadata,
                timestamp_consistency=timestamp_consistency,
                software_signatures=software_signatures,
                anomalies=anomalies,
                geo_location=geo_location,
                device_fingerprint=device_fingerprint
            )

        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return MetadataAnalysis(
                document_id=0,
                extracted_metadata={"error": str(e)},
                anomalies=[MetadataAnomaly(
                    anomaly_type="extraction_error",
                    description=f"Failed to extract metadata: {str(e)}",
                    severity=RiskLevel.HIGH,
                    affected_fields=["all"],
                    confidence=1.0
                )]
            )

    async def _extract_image_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata from image files using EXIF data."""
        metadata = {}
        
        try:
            # Use PIL for basic EXIF extraction
            image = Image.open(io.BytesIO(file_content))
            
            # Basic image properties
            metadata.update({
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height
            })
            
            # Extract EXIF data using PIL
            exif_data = image.getexif()
            if exif_data:
                exif_dict = {}
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    exif_dict[tag] = value
                metadata['exif_pil'] = exif_dict
                
                # Extract GPS data if available
                gps_info = exif_data.get_ifd(0x8825)  # GPS IFD
                if gps_info:
                    gps_dict = {}
                    for tag_id, value in gps_info.items():
                        tag = GPSTAGS.get(tag_id, tag_id)
                        gps_dict[tag] = value
                    metadata['gps_data'] = gps_dict

            # Use exifread for more detailed extraction
            file_content_copy = io.BytesIO(file_content)
            exif_tags = exifread.process_file(file_content_copy, details=True)
            
            if exif_tags:
                exif_detailed = {}
                for tag, value in exif_tags.items():
                    if tag not in ['JPEGThumbnail', 'TIFFThumbnail', 'Filename', 'EXIF MakerNote']:
                        exif_detailed[tag] = str(value)
                metadata['exif_detailed'] = exif_detailed

        except Exception as e:
            logger.error(f"Error extracting image metadata: {str(e)}")
            metadata['extraction_error'] = str(e)
            
        return metadata

    async def _extract_pdf_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata = {}
        
        try:
            pdf_reader = PdfReader(io.BytesIO(file_content))
            
            # Basic PDF properties
            metadata.update({
                'page_count': len(pdf_reader.pages),
                'is_encrypted': pdf_reader.is_encrypted
            })
            
            # Extract document information
            if pdf_reader.metadata:
                doc_info = {}
                for key, value in pdf_reader.metadata.items():
                    # Remove the leading '/' from PDF metadata keys
                    clean_key = key.lstrip('/')
                    doc_info[clean_key] = str(value) if value else None
                metadata['document_info'] = doc_info
                
                # Extract specific timestamps
                creation_date = pdf_reader.metadata.get('/CreationDate')
                modification_date = pdf_reader.metadata.get('/ModDate')
                
                if creation_date:
                    metadata['creation_date'] = str(creation_date)
                if modification_date:
                    metadata['modification_date'] = str(modification_date)
                    
                # Extract author and creator information
                author = pdf_reader.metadata.get('/Author')
                creator = pdf_reader.metadata.get('/Creator')
                producer = pdf_reader.metadata.get('/Producer')
                
                if author:
                    metadata['author'] = str(author)
                if creator:
                    metadata['creator'] = str(creator)
                if producer:
                    metadata['producer'] = str(producer)

            # Extract form fields if present
            if hasattr(pdf_reader, 'get_form_text_fields'):
                try:
                    form_fields = pdf_reader.get_form_text_fields()
                    if form_fields:
                        metadata['form_fields'] = form_fields
                except:
                    pass  # Form fields not available in this PDF

        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {str(e)}")
            metadata['extraction_error'] = str(e)
            
        return metadata

    async def _extract_docx_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata from DOCX files."""
        metadata = {}
        
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            
            # Core document properties
            core_props = doc.core_properties
            
            metadata.update({
                'author': core_props.author,
                'category': core_props.category,
                'comments': core_props.comments,
                'content_status': core_props.content_status,
                'created': core_props.created.isoformat() if core_props.created else None,
                'identifier': core_props.identifier,
                'keywords': core_props.keywords,
                'language': core_props.language,
                'last_modified_by': core_props.last_modified_by,
                'last_printed': core_props.last_printed.isoformat() if core_props.last_printed else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'revision': core_props.revision,
                'subject': core_props.subject,
                'title': core_props.title,
                'version': core_props.version
            })
            
            # Document statistics
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            
            # Count words and characters (basic implementation)
            word_count = 0
            char_count = 0
            for paragraph in doc.paragraphs:
                text = paragraph.text
                word_count += len(text.split())
                char_count += len(text)
                
            metadata.update({
                'paragraph_count': paragraphs,
                'table_count': tables,
                'word_count': word_count,
                'character_count': char_count
            })

        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
            metadata['extraction_error'] = str(e)
            
        return metadata

    async def _extract_xlsx_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata from XLSX files."""
        metadata = {}
        
        try:
            workbook = load_workbook(io.BytesIO(file_content))
            
            # Basic workbook properties
            metadata.update({
                'worksheet_count': len(workbook.worksheets),
                'worksheet_names': workbook.sheetnames
            })
            
            # Document properties
            props = workbook.properties
            if props:
                metadata.update({
                    'title': props.title,
                    'subject': props.subject,
                    'creator': props.creator,
                    'last_modified_by': props.lastModifiedBy,
                    'created': props.created.isoformat() if props.created else None,
                    'modified': props.modified.isoformat() if props.modified else None,
                    'description': props.description,
                    'keywords': props.keywords,
                    'category': props.category,
                    'version': props.version,
                    'revision': props.revision
                })
            
            # Calculate basic statistics
            total_cells = 0
            total_rows = 0
            for worksheet in workbook.worksheets:
                if worksheet.max_row:
                    total_rows += worksheet.max_row
                if worksheet.max_column and worksheet.max_row:
                    total_cells += worksheet.max_column * worksheet.max_row
                    
            metadata.update({
                'total_rows': total_rows,
                'estimated_total_cells': total_cells
            })

        except Exception as e:
            logger.error(f"Error extracting XLSX metadata: {str(e)}")
            metadata['extraction_error'] = str(e)
            
        return metadata

    def _analyze_timestamp_consistency(self, metadata: Dict[str, Any]) -> Optional[TimestampConsistency]:
        """Analyze timestamp consistency across metadata fields."""
        timestamps = []
        anomalies = []
        
        # Collect all timestamp fields
        timestamp_fields = [
            'creation_date', 'modification_date', 'created', 'modified',
            'last_printed', 'DateTime', 'DateTimeOriginal', 'DateTimeDigitized'
        ]
        
        # Extract timestamps from various metadata sources
        for field in timestamp_fields:
            if field in metadata and metadata[field]:
                try:
                    # Handle different timestamp formats
                    timestamp_str = str(metadata[field])
                    # This is a simplified parser - in production, you'd want more robust parsing
                    if 'T' in timestamp_str or '-' in timestamp_str:
                        timestamps.append((field, timestamp_str))
                except:
                    continue
        
        # Check for nested metadata (like EXIF)
        for nested_key in ['exif_detailed', 'exif_pil', 'document_info']:
            if nested_key in metadata and isinstance(metadata[nested_key], dict):
                for field in timestamp_fields:
                    if field in metadata[nested_key] and metadata[nested_key][field]:
                        timestamps.append((f"{nested_key}.{field}", str(metadata[nested_key][field])))
        
        if not timestamps:
            return None
            
        # Basic consistency checks
        is_consistent = True
        chronological_order = True
        
        # Check for obvious anomalies
        if len(timestamps) > 1:
            # Look for timestamps in the future
            current_time = datetime.now()
            for field, timestamp_str in timestamps:
                try:
                    # Simple future date check
                    if '2030' in timestamp_str or '2040' in timestamp_str:
                        anomalies.append(f"Future timestamp detected in {field}: {timestamp_str}")
                        is_consistent = False
                except:
                    continue
        
        return TimestampConsistency(
            is_consistent=is_consistent,
            anomalies=anomalies,
            chronological_order=chronological_order,
            time_gaps=[]  # Would implement gap analysis in production
        )

    def _detect_software_signatures(self, metadata: Dict[str, Any]) -> List[SoftwareSignature]:
        """Detect software signatures from metadata."""
        signatures = []
        
        # Common software signature patterns
        software_patterns = {
            'Adobe Photoshop': ['Adobe Photoshop', 'Photoshop'],
            'Microsoft Word': ['Microsoft Office Word', 'Word'],
            'Microsoft Excel': ['Microsoft Office Excel', 'Excel'],
            'GIMP': ['GIMP'],
            'Canon': ['Canon'],
            'Nikon': ['NIKON'],
            'iPhone': ['iPhone'],
            'Android': ['Android']
        }
        
        # Search through all metadata fields
        all_text = str(metadata).lower()
        
        for software_name, patterns in software_patterns.items():
            for pattern in patterns:
                if pattern.lower() in all_text:
                    # Calculate confidence based on how specific the match is
                    confidence = 0.8 if len(pattern) > 5 else 0.6
                    
                    signatures.append(SoftwareSignature(
                        software_name=software_name,
                        version=None,  # Would extract version in production
                        confidence=confidence,
                        signature_type="metadata_string",
                        detection_method="pattern_matching"
                    ))
                    break  # Avoid duplicate signatures for the same software
        
        # Check specific metadata fields for more accurate detection
        creator_fields = ['creator', 'producer', 'Creator', 'Producer', 'Software']
        for field in creator_fields:
            if field in metadata and metadata[field]:
                value = str(metadata[field])
                signatures.append(SoftwareSignature(
                    software_name=value,
                    version=None,
                    confidence=0.9,
                    signature_type="creator_field",
                    detection_method="direct_extraction"
                ))
        
        return signatures

    def _detect_metadata_anomalies(self, metadata: Dict[str, Any]) -> List[MetadataAnomaly]:
        """Detect anomalies in metadata that might indicate tampering."""
        anomalies = []
        
        # Check for extraction errors
        if 'extraction_error' in metadata:
            anomalies.append(MetadataAnomaly(
                anomaly_type="extraction_failure",
                description=f"Failed to extract metadata: {metadata['extraction_error']}",
                severity=RiskLevel.MEDIUM,
                affected_fields=["metadata_extraction"],
                confidence=1.0
            ))
        
        # Check for missing expected metadata
        if not any(key in metadata for key in ['creation_date', 'created', 'DateTime']):
            anomalies.append(MetadataAnomaly(
                anomaly_type="missing_timestamps",
                description="No creation timestamp found in metadata",
                severity=RiskLevel.LOW,
                affected_fields=["timestamps"],
                confidence=0.7
            ))
        
        # Check for suspicious software signatures
        if 'exif_detailed' in metadata:
            exif_data = metadata['exif_detailed']
            # Look for signs of metadata manipulation tools
            suspicious_software = ['exiftool', 'metadata editor', 'photo mechanic']
            for software in suspicious_software:
                if any(software.lower() in str(value).lower() for value in exif_data.values()):
                    anomalies.append(MetadataAnomaly(
                        anomaly_type="suspicious_software",
                        description=f"Detected metadata manipulation software: {software}",
                        severity=RiskLevel.MEDIUM,
                        affected_fields=["software_signatures"],
                        confidence=0.8
                    ))
        
        return anomalies

    def _extract_geo_location(self, metadata: Dict[str, Any]) -> Optional[GeoLocation]:
        """Extract geographic location from GPS metadata."""
        try:
            # Look for GPS data in EXIF
            if 'gps_data' in metadata:
                gps_data = metadata['gps_data']
                
                # Extract latitude and longitude
                lat_ref = gps_data.get('GPSLatitudeRef', 'N')
                lat_coords = gps_data.get('GPSLatitude')
                lon_ref = gps_data.get('GPSLongitudeRef', 'E')
                lon_coords = gps_data.get('GPSLongitude')
                
                if lat_coords and lon_coords:
                    # Convert GPS coordinates to decimal degrees
                    # This is a simplified conversion - production code would be more robust
                    try:
                        if isinstance(lat_coords, (list, tuple)) and len(lat_coords) >= 2:
                            lat = float(lat_coords[0]) + float(lat_coords[1]) / 60
                            if len(lat_coords) > 2:
                                lat += float(lat_coords[2]) / 3600
                            if lat_ref == 'S':
                                lat = -lat
                                
                            lon = float(lon_coords[0]) + float(lon_coords[1]) / 60
                            if len(lon_coords) > 2:
                                lon += float(lon_coords[2]) / 3600
                            if lon_ref == 'W':
                                lon = -lon
                                
                            # Extract altitude if available
                            altitude = None
                            if 'GPSAltitude' in gps_data:
                                altitude = float(gps_data['GPSAltitude'])
                                
                            return GeoLocation(
                                latitude=lat,
                                longitude=lon,
                                altitude=altitude
                            )
                    except (ValueError, TypeError, IndexError):
                        pass
                        
        except Exception as e:
            logger.error(f"Error extracting geo location: {str(e)}")
            
        return None

    def _extract_device_fingerprint(self, metadata: Dict[str, Any]) -> Optional[DeviceFingerprint]:
        """Extract device fingerprint information."""
        try:
            fingerprint_data = {}
            
            # Extract camera information from EXIF
            if 'exif_detailed' in metadata:
                exif_data = metadata['exif_detailed']
                
                # Camera make and model
                make = exif_data.get('Image Make', '').strip()
                model = exif_data.get('Image Model', '').strip()
                
                if make:
                    fingerprint_data['camera_make'] = make
                if model:
                    fingerprint_data['camera_model'] = model
                    
                # Lens information
                lens_info = exif_data.get('EXIF LensModel', '') or exif_data.get('EXIF LensMake', '')
                if lens_info:
                    fingerprint_data['lens_info'] = lens_info.strip()
                    
                # Software version
                software = exif_data.get('Image Software', '').strip()
                if software:
                    fingerprint_data['software_version'] = software
                    
                # Unique identifiers
                unique_ids = {}
                serial_fields = ['EXIF BodySerialNumber', 'EXIF LensSerialNumber', 'Image SerialNumber']
                for field in serial_fields:
                    if field in exif_data and exif_data[field]:
                        unique_ids[field] = str(exif_data[field]).strip()
                        
                if unique_ids:
                    fingerprint_data['unique_identifiers'] = unique_ids
            
            # Extract from PIL EXIF data as backup
            if 'exif_pil' in metadata and not fingerprint_data:
                exif_pil = metadata['exif_pil']
                
                if 'Make' in exif_pil:
                    fingerprint_data['camera_make'] = str(exif_pil['Make'])
                if 'Model' in exif_pil:
                    fingerprint_data['camera_model'] = str(exif_pil['Model'])
                if 'Software' in exif_pil:
                    fingerprint_data['software_version'] = str(exif_pil['Software'])
            
            if fingerprint_data:
                return DeviceFingerprint(**fingerprint_data)
                
        except Exception as e:
            logger.error(f"Error extracting device fingerprint: {str(e)}")
            
        return None